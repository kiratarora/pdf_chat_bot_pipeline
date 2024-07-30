import pdfplumber, re
from pypdf import PdfReader 
from docx import Document
import pandas as pd

'''
Class to load text and tables from a pdf file.
'''
class PDFLoader:

    '''
    Setup for the class
    Parameters:
        file_path (str): the path of the pdf file that needs to parsed
    '''
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path

    '''
    Funtion to use pdf plumber to optically read text.
    Returns:
        str: extracted text
    '''
    def extract_plumber_text(self) -> str:
        text = ""
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text(x_tolerance=3, x_tolerance_ratio=None, y_tolerance=3, layout=False, x_density=7.25, y_density=13, line_dir_render=None, char_dir_render=None) + " "
        return text
    
    '''
    Funtion to use pdf plumber to optically read tables.
    Returns:
        n-D List: extracted tables
    '''
    def extract_table(self) -> list[list]:
        tables = []
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                tables += page.extract_tables(table_settings={})
        return tables
    
    '''
    Funtion that uses pdf reader to extract text
    Returns:
        str: extracted text
    '''
    def extract_reader_text(self) -> str:
        reader = PdfReader(self.file_path) 
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text

    '''
    Function that decides the best tool to use for exraction based on pre-defined conditions
    Returns:
        str: extracted text
    '''
    def parse_data(self) -> str:
        data_reader = self.extract_reader_text()
        data_plumber = self.extract_plumber_text()
        # len (more), chars [#,>,\n] (less)
        def count_chars(input_string):
            return len(re.findall(r'[!@#$%^&*()_+{}\[\]:;<>,.?\\/-]', input_string))
        reader_chars, plumber_chars = count_chars(data_reader),count_chars(data_plumber)
        if abs(reader_chars-plumber_chars) <= 20:
            if len(data_plumber)<len(data_reader):
                return data_reader 
            else:
                return data_plumber
        if reader_chars > plumber_chars:
            return data_reader
        else:
            return data_plumber

'''
Class to load text and tables from a doc file.
'''          
class DOCLoader:
    
    '''
    Setup for the class
    Parameters:
        file_path (str): the path of the pdf file that needs to parsed
    '''
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.document = Document(file_path)
    
    '''
    Funtion that uses doc reader to extract text
    Returns:
        str: extracted text
    ''' 
    def extract_text(self) -> str:
        text_list = [para.text for para in self.document.paragraphs]
        text = ''
        for i in text_list:
            text += i + '\n'
        return text
    
    '''
    Funtion that uses doc reader to extract table
    Returns:
       n-D List: extracted tables
    ''' 
    def extract_table(self) -> list[list]:
        tables_data = []
        
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        return tables_data

'''
Class to load tables from a xlsx or a csv file.
'''   
class XLSLoader:
    
    '''
    Constructor for the class
    Parameters:
        file_path (str): the path of the pdf file that needs to parsed
    '''
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
    
    '''
    Function to load the csv and convert it to a n-D list
    Returns:
        n-D list: extracted tables
    '''
    def read_csv(self) -> list[list]:
        return pd.read_csv(self.file_path).values.tolist()
    
    '''
    Function to load a excel file and convert it to a n-D list
    Parameter:
        sheet_name (int): name of the excel sheet
    Returns:
        n-D list: extracted tables
    ''' 
    def read_excel(self, sheet_name: int = 0) -> list[list]:
        return pd.read_excel(self.file_path, sheet_name=sheet_name).values.tolist()

            